# 文件用途：专利 MD→DOCX 核心转换引擎。基于 markdown-it-py AST Token 流，调度专利排版函数生成符合 CNIPA 规范的 DOCX。
# 上游：MCP server.py 调用 build_patent_docx()；依赖 patent_style_utils.py 段落构建函数族
# 下游：输出 python-docx Document 对象，由调用方 .save() 落盘
# 核心路径：markdown-it-py 解析 → Token 遍历 → 语义分类 → patent_style_utils 段落构建

import re
import markdown_it
from docx import Document

from tools.patent_style_utils import (
    init_patent_document,
    add_section_title,
    add_sub_title,
    add_embodiment_title,
    add_invention_title,
    add_numbered_paragraph,
    add_claim_paragraph,
    add_step_paragraph,
    add_display_formula_placeholder,
    add_rich_paragraph,
)


# ============================================================================
# 语义分类 & 公式判定
# ============================================================================

def _classify_paragraph(text: str) -> str:
    """对顶层 paragraph 文本分类。返回: 'numbered'|'step'|'formula'|'normal'。"""
    stripped = text.strip()
    if not stripped:
        return 'empty'

    if re.match(r'^\[(\d{4})\]\s', stripped):
        return 'numbered'

    if re.match(r'^S\d+\.\s', stripped):
        return 'step'

    if _is_standalone_formula(stripped):
        return 'formula'

    return 'normal'


def _is_standalone_formula(text: str) -> bool:
    """判断行内容是否为独立公式（整行仅由 $...$ 组成且无非空白正文）。"""
    stripped = text.strip()
    if not stripped:
        return False
    return bool(re.match(r'^\$[^$]+\$$', stripped))


# ============================================================================
# 表格子解析器
# ============================================================================

def _parse_table_block(doc: Document, tokens: list, start_idx: int) -> int:
    """消费 table_open → ... → table_close 的连续 token 序列，构建 DOCX 表格。返回下一个未消费的 token 索引。"""
    i = start_idx + 1
    rows_data = []
    current_row = []
    in_thead = False

    while i < len(tokens) and tokens[i].type != 'table_close':
        t = tokens[i]

        if t.type == 'thead_open':
            in_thead = True
        elif t.type == 'thead_close':
            in_thead = False
        elif t.type == 'tbody_open':
            in_thead = False
        elif t.type == 'tr_open':
            current_row = []
        elif t.type == 'tr_close':
            if current_row:
                rows_data.append((in_thead, current_row))
        elif t.type in ('td_open', 'th_open'):
            inline = tokens[i + 1]
            cell_text = inline.content.strip() if inline.content else ''
            current_row.append(cell_text)
            i += 2  # 跳过 inline 和 td_close/th_close
            continue

        i += 1

    if not rows_data:
        return i + 1

    num_cols = max(len(r[1]) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = 1  # CENTER

    from docx.shared import Pt
    from tools.patent_style_utils import set_run_font
    from docx.oxml.ns import qn

    for row_idx, (is_header, cells) in enumerate(rows_data):
        for col_idx, cell_text in enumerate(cells):
            if col_idx >= num_cols:
                break
            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            p.alignment = 1  # CENTER
            run = p.add_run(cell_text)
            if is_header:
                run.bold = True
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'D9E2F3',
                    qn('w:val'): 'clear'
                })
                shading.append(shd)
            set_run_font(run, size=Pt(9))

    return i + 1


# ============================================================================
# ordered_list 权利要求块处理
# ============================================================================

def _parse_ordered_list_block(doc: Document, tokens: list, start_idx: int) -> int:
    """消费 ordered_list_open → list_item_open → ... → ordered_list_close。每个 list_item 为一项权利要求。返回下一个未消费的 token 索引。"""
    i = start_idx + 1
    claim_num = 1
    start_attr = tokens[start_idx].attrGet('start')
    if start_attr:
        claim_num = int(start_attr)

    while i < len(tokens) and tokens[i].type != 'ordered_list_close':
        if tokens[i].type != 'list_item_open':
            i += 1
            continue

        body_parts = []
        i += 1
        while i < len(tokens) and tokens[i].type != 'list_item_close':
            if tokens[i].type == 'paragraph_open':
                inline = tokens[i + 1]
                text = inline.content.strip()
                if text:
                    body_parts.append(text)
                i += 3
            else:
                i += 1

        body = '; '.join(body_parts) if body_parts else ''
        add_claim_paragraph(doc, claim_num, body)
        claim_num += 1

    return i + 1


# ============================================================================
# 主入口
# ============================================================================

def build_patent_docx(md_text: str) -> Document:
    """专利 MD → CNIPA 合规 DOCX 主入口。返回 Document 对象，调用方负责 .save()。"""
    md = markdown_it.MarkdownIt('gfm-like', {'breaks': False})  # gfm-like 启用表格解析 (commonmark 不含表格)
    tokens = md.parse(md_text)
    doc = init_patent_document()

    i = 0
    prev_was_h1_description = False  # 仅 # 说明书 后的首个 paragraph 为发明名称
    has_drawing_title = False        # 检测 MD 是否已有 # 说明书附图

    while i < len(tokens):
        t = tokens[i]

        # ---- heading 处理 ----
        if t.type == 'heading_open':
            level = int(t.tag[1])
            text = tokens[i + 1].content.strip()

            if level == 1:
                add_section_title(doc, text)
                prev_was_h1_description = (text == '说明书')
                if text == '说明书附图':
                    has_drawing_title = True
            elif level == 2:
                add_sub_title(doc, text)
            elif level == 3:
                add_embodiment_title(doc, text)

            i += 3
            continue

        # ---- ordered_list (权利要求) ----
        if t.type == 'ordered_list_open':
            i = _parse_ordered_list_block(doc, tokens, i)
            continue

        # ---- 表格 ----
        if t.type == 'table_open':
            i = _parse_table_block(doc, tokens, i)
            continue

        # ---- 顶层 paragraph ----
        if t.type == 'paragraph_open':
            inline = tokens[i + 1]
            raw = inline.content.strip()

            # 发明名称: # 说明书 后的第一个非空 paragraph
            if prev_was_h1_description and raw:
                add_invention_title(doc, raw)
                prev_was_h1_description = False
                i += 3
                continue
            prev_was_h1_description = False

            # 空段落 → 保留段落间距
            if not raw:
                doc.add_paragraph()
                i += 3
                continue

            tag = _classify_paragraph(raw)

            if tag == 'numbered':
                m = re.match(r'^\[(\d{4})\]\s(.*)', raw)
                if m:
                    add_numbered_paragraph(doc, m.group(1), m.group(2))
                else:
                    add_rich_paragraph(doc, inline)
            elif tag == 'step':
                m = re.match(r'^(S\d+\.)\s(.*)', raw)
                if m:
                    add_step_paragraph(doc, m.group(1), m.group(2))
                else:
                    add_rich_paragraph(doc, inline)
            elif tag == 'formula':
                add_display_formula_placeholder(doc, raw)
            else:
                add_rich_paragraph(doc, inline)

            i += 3
            continue

        i += 1

    # 末尾兜底：若 MD 未含说明书附图标题，自动追加
    if not has_drawing_title:
        add_section_title(doc, "说明书附图")

    # 纯净着陆区：阻断 H1 边框/字号向下级联，供用户安全粘贴附图
    p_landing = doc.add_paragraph()
    p_landing.style = doc.styles['Normal']

    return doc
