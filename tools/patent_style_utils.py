# 文件用途：专利 MD→DOCX 底层工具层。提供字体引擎、页面初始化、各类专利段落构建函数。
# 上游：patent_md_to_docx.py 调用本模块的所有段落构建函数
# 下游：通过 python-docx 生成符合 CNIPA 规范的 DOCX 段落和 Run 对象
# 核心职责：封装所有 Word OOXML 底层操作（字体/页边距/缩进/底纹），上层无需触碰 XML

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
import re


# ============================================================================
# 字体引擎 & 页面初始化
# ============================================================================

def set_run_font(run, cn_font='宋体', en_font='Times New Roman',
                 size=Pt(12), bold=False):
    """一步到位设置中英文混排字体+字号+加粗。含 w:hint='eastAsia' 防御性渲染提示。"""
    run.font.name = en_font
    run.font.size = size
    run.font.bold = bold

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:hint'), 'eastAsia')


def init_patent_document() -> Document:
    """创建符合 CNIPA 规范的空白专利文档。页边距: 左25/上25/右15/下15mm。"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)

    return doc


# ============================================================================
# 段落构建函数族
# ============================================================================

def add_section_title(doc: Document, text: str):
    """五大书一级标题。3空格字距 + 仅加粗下边框 + 楷体 18pt 居中。"""
    # 暴力清洗：去图号 → 清所有空白 → 3空格重排
    clean = re.sub(r'[\(（]图\d+[\)）]', '', text).strip()
    clean = re.sub(r'\s+', '', clean)
    spaced = '   '.join(list(clean))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.page_break_before = True  # 原生分页，替代硬分页符

    run = p.add_run(spaced)
    set_run_font(run, '楷体_GB2312', 'Times New Roman', Pt(18), bold=False)

    # OOXML: 仅加粗下边框 (w:sz=16 → 2pt)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '16')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_sub_title(doc: Document, text: str):
    """说明书子标题。宋体 12pt 加粗。摘要附图居中，其余左对齐顶格。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    if '摘要附图' in text:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = 0

    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(12), bold=True)
    return p


def add_embodiment_title(doc: Document, text: str):
    """实施例标题。宋体 12pt 加粗左对齐顶格。"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = 0
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(12), bold=True)
    return p


def add_invention_title(doc: Document, text: str):
    """发明名称。居中加粗 16pt。CNIPA 要求: 发明名称与正文之间空一行。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(16), bold=True)
    doc.add_paragraph()  # CNIPA 规范: 发明名称与正文之间空一行
    return p


def _make_body_paragraph(doc: Document) -> 'Paragraph':
    """创建带标准正文格式的空白段落 (首行缩进 Pt(24), 行距 1.5)。"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_numbered_paragraph(doc: Document, number: str, body: str):
    """说明书编号段落。编号加粗 + 正文缩进 Pt(24)。"""
    p = _make_body_paragraph(doc)
    bold_run = p.add_run(f'[{number}] ')
    set_run_font(bold_run, size=Pt(12), bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=Pt(12))
    return p


def add_claim_paragraph(doc: Document, number: int, body: str):
    """权利要求项。无首行缩进，编号加粗。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    bold_run = p.add_run(f'{number}. ')
    set_run_font(bold_run, size=Pt(12), bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=Pt(12))
    return p


def add_step_paragraph(doc: Document, step_label: str, body: str):
    """方法步骤子项 (S1. xxx)。编号加粗 + 正文缩进 Pt(24)。"""
    p = _make_body_paragraph(doc)
    bold_run = p.add_run(f'{step_label} ')
    set_run_font(bold_run, size=Pt(12), bold=True)
    body_run = p.add_run(body)
    set_run_font(body_run, size=Pt(12))
    return p


def add_display_formula_placeholder(doc: Document, latex_text: str):
    """块级公式占位。居中纯黑文本，保留 $ 分隔符供 MathType 宏捕获。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(latex_text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(12))
    return p


def add_rich_paragraph(doc: Document, inline_token, cn_font='宋体',
                       en_font='Times New Roman', size=Pt(12)):
    """
    处理包含行内元素的段落。
    遍历 inline.children 映射 strong_open/em_open/text 到 DOCX Run。
    T2: $...$ 公式不再拆分，作为普通文本单 Run 原样写入以保证 MathType 宏兼容。
    """
    p = _make_body_paragraph(doc)
    children = inline_token.children if inline_token.children else []

    in_strong = False
    in_em = False

    for child in children:
        if child.type == 'strong_open':
            in_strong = True
            continue
        if child.type == 'strong_close':
            in_strong = False
            continue
        if child.type == 'em_open':
            in_em = True
            continue
        if child.type == 'em_close':
            in_em = False
            continue
        if child.type != 'text':
            continue

        txt = child.content
        if not txt.strip():
            continue

        run = p.add_run(txt)
        run.bold = in_strong
        run.italic = in_em
        set_run_font(run, cn_font, en_font, size)

    return p
