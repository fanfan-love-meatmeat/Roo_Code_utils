import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_run_font(run, size=Pt(10.5), bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

def set_heading_font(paragraph, level):
    size_map = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(11)}
    for run in paragraph.runs:
        set_run_font(run, size=size_map.get(level, Pt(11)), bold=True)

def parse_md_to_docx(md_path: str, docx_path: str = None) -> str:
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Markdown 文件不存在: {md_path}")

    # 智能推导输出目录：默认 -> 输入文件同级的相同名字.docx
    if docx_path is None:
        docx_path = os.path.splitext(md_path)[0] + ".docx"

    os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title_text = stripped.lstrip("#").strip()
            title_text = re.sub(r"\*\*(.*?)\*\*", r"\1", title_text)
            level = min(max(level, 1), 4) # 限制 1-4 级
            p = doc.add_heading(title_text, level=level)
            set_heading_font(p, level)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) < 2:
                continue

            header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
            data_rows = []
            for tl in table_lines[2:]: # 跳过分隔符行
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                data_rows.append(cells)

            num_cols = len(header_cells)
            table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, cell_text in enumerate(header_cells):
                cell = table.rows[0].cells[j]
                cell_text_clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                cell_text_clean = re.sub(r"\[.*?\]\(.*?\)", "", cell_text_clean)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cell_text_clean)
                set_run_font(run, size=Pt(9), bold=True)
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): "D9E2F3", qn("w:val"): "clear"
                })
                shading.append(shading_elem)

            for row_idx, row_data in enumerate(data_rows):
                for j in range(min(len(row_data), num_cols)):
                    cell = table.rows[row_idx + 1].cells[j]
                    cell_text = row_data[j]
                    cell_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                    cell_text = cell_text.replace("✅", "√").replace("✓", "√")
                    cell_text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", cell_text)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(cell_text)
                    set_run_font(run, size=Pt(9))
            
            doc.add_paragraph()
            continue

        if stripped.startswith(">"):
            quote_text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped.lstrip("> ").strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(quote_text)
            set_run_font(run, size=Pt(9))
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        if stripped.startswith("- "):
            item_text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped[2:].strip())
            p = doc.add_paragraph(item_text, style="List Bullet")
            for run in p.runs:
                set_run_font(run, size=Pt(10))
            i += 1
            continue

        text = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
        text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
        p = doc.add_paragraph(text)
        for run in p.runs:
            set_run_font(run, size=Pt(10.5))
        i += 1

    doc.save(docx_path)
    print(f"[md_to_docx] 转换完成: {docx_path}")
    return docx_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python utils/md_to_docx.py <md文件路径> [输出docx路径]")
        sys.exit(1)

    md = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    parse_md_to_docx(md, out)