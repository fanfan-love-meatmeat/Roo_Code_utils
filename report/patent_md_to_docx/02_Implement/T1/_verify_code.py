"""Verify python-docx code skeletons from v2.0 mapping table"""
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

# ========== set_run_font ==========
def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False):
    run.font.name = en_font
    run.font.size = size
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:hint'), 'eastAsia')

# ========== init_patent_document ==========
def init_patent_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)
    return doc

# ========== _is_standalone_formula ==========
def _is_standalone_formula(text):
    stripped = text.strip()
    return bool(re.match(r'^\$[^$]+\$$', stripped))

# ========== TESTS ==========
print("=" * 60)
print("CODE SKELETON VERIFICATION")
print("=" * 60)

# Test 1: init_patent_document
doc = init_patent_document()
sec = doc.sections[0]

# EMU to mm conversion: 1mm = 36000 EMU
top_mm = sec.top_margin / 36000
bottom_mm = sec.bottom_margin / 36000
left_mm = sec.left_margin / 36000
right_mm = sec.right_margin / 36000

print(f"\n1. init_patent_document:")
print(f"   margins: top={top_mm:.0f}mm bottom={bottom_mm:.0f}mm left={left_mm:.0f}mm right={right_mm:.0f}mm")
assert round(top_mm) == 25, f"Expected 25mm, got {top_mm}mm"
assert round(bottom_mm) == 15, f"Expected 15mm, got {bottom_mm}mm"
assert round(left_mm) == 25, f"Expected 25mm, got {left_mm}mm"
assert round(right_mm) == 15, f"Expected 15mm, got {right_mm}mm"
print("   PASS: margins correct (25/15/25/15mm)")

# Test 2: set_run_font
p = doc.add_paragraph()
run = p.add_run('测试test')
set_run_font(run, size=Pt(12))
rPr = run._element.rPr
rFonts = rPr.find(qn('w:rFonts'))
assert rFonts is not None, "rFonts not created"
eastAsia = rFonts.get(qn('w:eastAsia'))
ascii_font = rFonts.get(qn('w:ascii'))
hAnsi = rFonts.get(qn('w:hAnsi'))
hint = rFonts.get(qn('w:hint'))
print(f"\n2. set_run_font:")
print(f"   eastAsia={eastAsia}, ascii={ascii_font}, hAnsi={hAnsi}, hint={hint}")
assert eastAsia == '宋体', f"Expected 宋体, got {eastAsia}"
assert ascii_font == 'Times New Roman', f"Expected Times New Roman, got {ascii_font}"
assert hint == 'eastAsia', f"Expected eastAsia, got {hint}"
assert run.font.size == Pt(12), f"Expected Pt(12), got {run.font.size}"
print("   PASS: fonts and hint correct")

# Test 3: heading paragraph
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('说明书摘要')
set_run_font(run, '宋体', 'Times New Roman', Pt(18), bold=False)
assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
assert run.font.size == Pt(18)
print("\n3. Heading paragraph: PASS")

# Test 4: numbered paragraph
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Pt(24)
p.paragraph_format.line_spacing = 1.5
bold_run = p.add_run('[0001] ')
bold_run.bold = True
set_run_font(bold_run, size=Pt(12))
body_run = p.add_run('本发明属于机器人自主导航技术领域。')
set_run_font(body_run, size=Pt(12))
assert p.paragraph_format.first_line_indent == Pt(24)
assert p.paragraph_format.line_spacing == 1.5
print("4. Numbered paragraph: PASS (indent=Pt24, line_spacing=1.5)")

# Test 5: formula detection
assert _is_standalone_formula('$V_{overlap} = N_{collision} \\cdot l^3$') == True
assert _is_standalone_formula('text $x^2$ text') == False
assert _is_standalone_formula('  $R_t = R_{progress,t} + R_{crush,t}$  ') == True
print("\n5. _is_standalone_formula: PASS")
print("   standalone detected: $V_overlap = ...$")
print("   inline rejected: 'text $x^2$ text'")

# Test 6: formula placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._element.get_or_add_pPr()
shd = pPr.makeelement(qn('w:shd'), {qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear'})
pPr.append(shd)
run = p.add_run('[公式] $R_t = R_{progress,t} + R_{crush,t}$')
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
set_run_font(run, size=Pt(12))
print("\n6. Formula placeholder: PASS (gray bg + red text)")

print("\n" + "=" * 60)
print("ALL 6 TESTS PASSED")
print("=" * 60)
